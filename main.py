from docx import Document
from docx.shared import RGBColor, Pt
from docx.shared import Cm, Inches, Pt
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
import tkinter.filedialog
from tkinter import messagebox



inFilmPath = path=tkinter.filedialog.askopenfilename()

eFont = 'Times New Roman'
cFont = '华文宋体'
margin = 0.5

document = Document()

cStyle = document.styles.add_style('cStyle', WD_STYLE_TYPE.CHARACTER)
cStyle.font.name = cFont
document.styles['cStyle']._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文宋体')

fin = open(inFilmPath, "r", encoding='utf-16')

firstLine = None
read = True
while True:
    if read == True:
        line = fin.readline().strip()

    if line == '':
        break

    if line[0] == '+':
        line = line[1:]

        paragraph = document.add_paragraph()

        paragraph_format = paragraph.paragraph_format
        tab_stops = paragraph_format.tab_stops
        tab_stop = tab_stops.add_tab_stop(Cm(5))

        paragraph.space_after = Pt(12)
        run = paragraph.add_run()
        run.text = line
        run.font.size = Pt(14)
        run.bold = True
        firstLine = True

    if line[0] == '#':
        line = line[1:]
        if firstLine == False:
            paragraph.add_run('\n')

        run = paragraph.add_run()
        run.text = '\t' + line
        run.font.size = Pt(10)
        run.font.name = eFont
        run.style = "cStyle"
        run.bold = False
        firstLine = False

section = document.sections[0]
section.top_margin = Cm(margin)
section.bottom_margin = Cm(margin)
section.left_margin = Cm(margin)
section.right_margin = Cm(margin)

fin.close()
document.save(inFilmPath+".docx")
messagebox.showinfo(title='导出完成', message='文档已导出\n导出文档地址：'+inFilmPath+".docx")